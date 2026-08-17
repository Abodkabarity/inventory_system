from __future__ import annotations

import argparse
import hashlib
import mimetypes
import os
import re
import socket
import sys
import uuid
from datetime import UTC, datetime
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Any, Iterable

VENDOR_DIR = Path(__file__).resolve().parent / ".vendor"
if VENDOR_DIR.is_dir() and str(VENDOR_DIR) not in sys.path:
    sys.path.insert(0, str(VENDOR_DIR))

import pdfplumber
from docx import Document
from openpyxl import load_workbook

if TYPE_CHECKING:
    from supabase import Client
else:
    Client = Any

try:
    import ftfy
except ImportError:  # The worker still preserves text if optional mojibake repair is unavailable.
    ftfy = None


SUPPORTED = {".pdf", ".docx", ".xlsx", ".xlsb"}
MIME_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xlsb": "application/vnd.ms-excel.sheet.binary.macroEnabled.12",
}


@dataclass(slots=True)
class ExtractedBlock:
    text: str
    raw: str
    content_type: str = "paragraph"
    extraction_method: str = "native"
    page_from: int | None = None
    page_to: int | None = None
    sheet_name: str | None = None
    row_from: int | None = None
    row_to: int | None = None
    section_title: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


def normalize_text(value: str) -> str:
    value = ftfy.fix_text(value or "") if ftfy else (value or "")
    # Some embedded PDF fonts expose the "ti" ligature as U+FFFD. Repair only
    # high-confidence policy headings; unknown replacement characters remain
    # visible so the post-index validator can reject the document for review.
    value = re.sub(r"\bAdministra\ufffdon\b", "Administration", value, flags=re.IGNORECASE)
    value = re.sub(r"\bDura\ufffdon\b", "Duration", value, flags=re.IGNORECASE)
    value = value.replace("\u00a0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r" *\n *", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def token_estimate(value: str) -> int:
    return max(1, len(re.findall(r"\w+|[^\w\s]", value, flags=re.UNICODE)))


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def content_hash(value: str) -> str:
    return hashlib.sha256(normalize_text(value).encode("utf-8")).hexdigest()


def utc_now() -> str:
    return datetime.now(UTC).isoformat()


def table_to_text(table: list[list[str | None]]) -> str:
    rows: list[str] = []
    for row in table:
        cells = [normalize_text(str(cell or "")) for cell in row]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def normalize_field_name(value: str) -> str:
    value = normalize_text(value).lower()
    value = re.sub(r"[^a-z0-9\u0600-\u06ff]+", "_", value, flags=re.UNICODE)
    return value.strip("_") or "field"


def normalize_entity_name(value: str) -> str:
    value = normalize_text(value).casefold()
    value = re.sub(r"[^\w\u0600-\u06ff]+", " ", value, flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip()


PROSE_MEDICATION_STRENGTH = re.compile(
    r"\b([A-Z][A-Za-z0-9-]{2,}(?:\s+[A-Z][A-Za-z0-9-]{2,}){0,2})"
    r"\s*:?\s+(\d+(?:\.\d+)?)\s*((?i:mg|mcg|g|gram(?:s)?|ml))\b",
)
ENTITY_CANDIDATE_STOPWORDS = {
    "age", "and", "after", "before", "daily", "dose", "exceeding", "for",
    "initial", "initiate", "maximum", "minimum", "monthly", "months", "need",
    "recommended", "reaching", "than", "the", "then", "three", "two", "until",
}


def extract_entity_aliases(
    chunks: list[ExtractedBlock],
    *,
    document_id: str,
    document_title: str,
) -> list[dict[str, Any]]:
    """Create a query-time entity registry from structured rows and prose strengths."""
    aliases: dict[tuple[str, str], dict[str, Any]] = {}

    def register(
        name: str,
        discovery: str,
        strength: str | None = None,
        entity_type: str = "medication",
    ) -> None:
        normalized = normalize_entity_name(name)
        words = normalized.split()
        if (
            not normalized
            or any(word in ENTITY_CANDIDATE_STOPWORDS for word in words)
            or all(len(word) <= 2 for word in words)
        ):
            return
        canonical = name.title() if name.isupper() else normalize_text(name)
        metadata: dict[str, Any] = {
            "document_id": document_id,
            "document_title": document_title,
            "discovery": discovery,
        }
        if strength:
            metadata["strength"] = strength
        aliases[(entity_type, normalized)] = {
            "entity_type": entity_type,
            "canonical_name": canonical,
            "alias": canonical,
            "normalized_alias": normalized,
            "language": "und",
            "metadata": metadata,
        }

    medication_name_suffix = re.compile(
        r"(?:mab|zumab|nib|gepant|pent|glutide|gliptin|statin|stim|setron|cept|afil)$",
        flags=re.IGNORECASE,
    )
    for chunk in chunks:
        if chunk.metadata.get("entity_type") == "medication" and chunk.metadata.get("entity_name"):
            register(str(chunk.metadata["entity_name"]), "structured_chunk")
        for match in PROSE_MEDICATION_STRENGTH.finditer(chunk.text):
            register(match.group(1), "prose_strength_mention", f"{match.group(2)} {match.group(3).lower()}")
        for token in re.findall(r"\b[A-Z][A-Za-z0-9-]{5,}\b", chunk.text):
            if medication_name_suffix.search(token):
                register(token, "medicine_suffix_mention")

    # The title is authoritative document metadata and remains useful when a
    # policy discusses a therapy without printing a dose next to its name.
    # Register acronym-like classes and medicine-looking title tokens without
    # treating generic title words as clinical entities.
    title_tokens = re.findall(r"[A-Za-z][A-Za-z0-9-]{2,}", document_title)
    title_stopwords = ENTITY_CANDIDATE_STOPWORDS | {
        "adjudication", "coverage", "guideline", "overview", "policy", "rule",
        "summary", "updated", "therapy", "therapies", "under", "daman",
    }
    for token in title_tokens:
        normalized = normalize_entity_name(token)
        if normalized in title_stopwords:
            continue
        if token.isupper() and 3 <= len(token) <= 10:
            register(token, "document_title_acronym", entity_type="therapy_class")
        elif medication_name_suffix.search(token):
            register(token, "document_title_medication")
    return list(aliases.values())


def document_profile(blocks: list[ExtractedBlock], *, title: str) -> dict[str, Any]:
    """Produce document-level structure metadata from trusted extraction output."""
    sections: list[dict[str, Any]] = []
    seen_paths: set[str] = set()
    for block in blocks:
        if not block.section_title:
            continue
        section_path = normalize_text(str(block.metadata.get("section_path") or block.section_title))
        if not section_path or section_path in seen_paths:
            continue
        seen_paths.add(section_path)
        sections.append({
            "title": normalize_text(str(block.section_title)),
            "path": section_path,
            "level": int(block.metadata.get("heading_level") or 1),
        })
    section_titles = [section["title"] for section in sections]
    table_rows = [block for block in blocks if block.content_type == "table_row"]
    field_names = sorted({
        key for block in table_rows
        for key in (block.metadata.get("fields") or {}).keys()
    })
    return {
        "title": title,
        "section_titles": section_titles,
        "sections": sections,
        "section_count": len(section_titles),
        "table_row_count": len(table_rows),
        "list_count": sum(block.content_type == "list" for block in blocks),
        "form_field_names": field_names,
        "content_types": sorted({block.content_type for block in blocks}),
    }


def extract_document_catalog(
    chunks: list[ExtractedBlock],
    *,
    document_id: str,
    document_title: str,
    aliases: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build a generic entity catalog using only trusted document structures."""
    catalog: dict[tuple[str, str], dict[str, Any]] = {}

    def register(name: str, entity_type: str, source: str, metadata: dict[str, Any] | None = None) -> None:
        canonical = normalize_text(name)
        normalized = normalize_entity_name(canonical)
        if not canonical or len(normalized) < 2:
            return
        key = (entity_type, normalized)
        catalog[key] = {
            "document_id": document_id,
            "entity_type": entity_type,
            "canonical_name": canonical,
            "normalized_entity": normalized,
            "aliases": [canonical],
            "source": source,
            "confidence": 1,
            "metadata": metadata or {},
        }

    for alias in aliases:
        entity_type = str(alias["entity_type"])
        register(
            str(alias["canonical_name"]),
            entity_type if entity_type in {"medication", "ingredient", "therapy_class"} else "policy_term",
            str((alias.get("metadata") or {}).get("discovery") or "document_alias"),
        )
    for chunk in chunks:
        if chunk.section_title:
            register(chunk.section_title, "policy_term", "section_heading", {"content_type": chunk.content_type})
        for field_name, field_value in (chunk.metadata.get("fields") or {}).items():
            entity_type = entity_type_for_header(normalize_field_name(str(field_name)))
            if not entity_type or not field_value:
                continue
            # A table cell may list several approved items; preserve each as a
            # separate catalog record without making assumptions about names.
            for value in re.split(r"\s*(?:,|;|\n|/|\band\b|\bor\b)\s*", str(field_value), flags=re.IGNORECASE):
                value = normalize_text(value)
                if len(normalize_entity_name(value)) >= 2:
                    register(value, entity_type, "structured_field", {"field": field_name})
        # Label/value prose is a trusted structure even outside tables.
        for label, value in re.findall(
            r"\b(medication|medicine|drug|brand|generic(?: name)?|active ingredient|diagnosis|condition|lab(?:oratory)? test|procedure)\s*:\s*([^\n.;]{2,100})",
            chunk.text,
            flags=re.IGNORECASE,
        ):
            entity_type = entity_type_for_header(normalize_field_name(label))
            if entity_type:
                register(value, entity_type, "labeled_prose", {"label": label})
    return list(catalog.values())


def document_health_report(
    chunks: list[ExtractedBlock],
    profile: dict[str, Any],
    catalog: list[dict[str, Any]],
) -> dict[str, Any]:
    """Fail closed before READY: verify structure and exact catalog term retrieval."""
    searchable = "\n".join(normalize_entity_name(chunk.text) for chunk in chunks)
    important = [entry for entry in catalog if entry["entity_type"] != "policy_term"]
    probes = important or [
        {"canonical_name": title}
        for title in profile.get("section_titles", [])[:8]
    ]
    failed_terms = [
        str(entry["canonical_name"])
        for entry in probes
        if normalize_entity_name(str(entry["canonical_name"])) not in searchable
    ]
    checks = {
        "has_chunks": bool(chunks),
        "has_structure": bool(profile.get("section_count") or profile.get("table_row_count")),
        "has_catalog": bool(catalog),
        "exact_terms_retrievable": not failed_terms,
    }
    return {
        "status": "verified" if all(checks.values()) else "failed",
        "checks": checks,
        "chunk_count": len(chunks),
        "catalog_entity_count": len(catalog),
        "failed_exact_terms": failed_terms[:20],
    }


ENTITY_HEADERS = {
    "medication": "medication",
    "medicine": "medication",
    "drug": "medication",
    "drug_name": "medication",
    "medication_name": "medication",
    "medicine_name": "medication",
    "brand": "medication",
    "brand_name": "medication",
    "generic_name": "medication",
    "active_ingredient": "medication",
    "ingredient": "medication",
    "product": "product",
    "item": "product",
    "item_name": "product",
}

# Classification is derived from document labels/columns, never from a known
# medicine list. This lets tomorrow's "Drug XYZ" become a searchable entity
# as soon as it appears under a trusted Medication/Brand/Generic field.
CATALOG_HEADER_TYPES = {
    "medication": "medication", "medicine": "medication", "drug": "medication",
    "brand": "brand", "brand_name": "brand", "generic": "generic",
    "generic_name": "generic", "active_ingredient": "ingredient",
    "ingredient": "ingredient", "therapy_class": "therapy_class",
    "drug_class": "therapy_class", "therapeutic_class": "therapy_class",
    "diagnosis": "diagnosis", "condition": "diagnosis", "indication": "diagnosis",
    "disease": "diagnosis", "laboratory": "lab", "lab": "lab", "test": "lab",
    "procedure": "procedure", "specialty": "policy_term", "prescriber": "policy_term",
}


def entity_type_for_header(key: str) -> str | None:
    if key in CATALOG_HEADER_TYPES:
        return CATALOG_HEADER_TYPES[key]
    tokens = set(key.split("_"))
    if {"medication", "medicine", "drug"} & tokens:
        return "medication"
    if {"brand", "generic", "ingredient"} & tokens:
        return "medication"
    if {"class", "therapy"} <= tokens:
        return "therapy_class"
    if {"diagnosis", "condition", "indication", "disease"} & tokens:
        return "diagnosis"
    if {"lab", "laboratory", "test"} & tokens:
        return "lab"
    if "procedure" in tokens:
        return "procedure"
    return None


def table_row_blocks(
    table: list[list[str | None]],
    *,
    extraction_method: str,
    section_title: str | None,
    table_index: int,
    section_path: str | None = None,
    page_number: int | None = None,
    sheet_name: str | None = None,
    source_file: str | None = None,
) -> list[ExtractedBlock]:
    """Convert a table into independent logical records without cross-row overlap."""
    normalized_rows = [
        [normalize_text(str(cell or "")) for cell in row]
        for row in table
    ]
    header_index = next((index for index, row in enumerate(normalized_rows) if any(row)), None)
    if header_index is None:
        return []

    headers = normalized_rows[header_index]
    width = len(headers)

    # PDF extractors frequently expose a visually wrapped one-column table as
    # several physical rows and mistakenly promote the first fragment to the
    # header. Keep that visual block together so a dose/criterion is never
    # split between unrelated searchable chunks.
    if width == 1:
        fragments = [row[0] for row in normalized_rows[header_index:] if row and row[0]]
        if not fragments:
            return []
        joined = normalize_text(" ".join(fragments))
        start = header_index + 1
        end = header_index + len(fragments)
        return [ExtractedBlock(
            text=joined,
            raw="\n".join(fragments),
            content_type="table_row",
            extraction_method=extraction_method,
            page_from=page_number,
            page_to=page_number,
            sheet_name=sheet_name,
            row_from=start,
            row_to=end,
            section_title=section_title,
            metadata={
                "table_index": table_index,
                "section_path": section_path or section_title,
                "table_title": section_title or fragments[0][:160],
                "headers": [],
                "fields": {"text": joined},
                "table_row_from": start,
                "table_row_to": end,
                "logical_row_group": f"{page_number or sheet_name or 'document'}:{table_index}:{start}-{end}",
                "merged_wrapped_rows": True,
                "source_file": source_file,
                "zero_cross_row_overlap": True,
            },
        )]
    keys: list[str] = []
    used_keys: dict[str, int] = {}
    for column, header in enumerate(headers):
        base = normalize_field_name(header or f"Column {column + 1}")
        used_keys[base] = used_keys.get(base, 0) + 1
        keys.append(base if used_keys[base] == 1 else f"{base}_{used_keys[base]}")

    entity_column = next((i for i, key in enumerate(keys) if entity_type_for_header(key) == "medication"), None)
    entity_type = "medication" if entity_column is not None else None
    table_title = " / ".join(header for header in headers if header) or section_title or "Table"
    logical_rows: list[tuple[int, int, list[str], list[list[str]]]] = []
    current_cells: list[str] | None = None
    current_raw_rows: list[list[str]] = []
    current_start = 0
    current_end = 0

    def flush() -> None:
        nonlocal current_cells, current_raw_rows, current_start, current_end
        if current_cells is not None and any(current_cells):
            logical_rows.append((current_start, current_end, current_cells, current_raw_rows))
        current_cells = None
        current_raw_rows = []

    for physical_index, original_row in enumerate(normalized_rows[header_index + 1 :], start=header_index + 2):
        row = (original_row + [""] * width)[:width]
        if not any(row):
            continue
        is_continuation = entity_column is not None and not row[entity_column] and current_cells is not None
        if not is_continuation:
            flush()
            current_cells = row.copy()
            current_raw_rows = [row.copy()]
            current_start = physical_index
            current_end = physical_index
            if entity_column is None:
                flush()
            continue
        for column, value in enumerate(row):
            if value:
                current_cells[column] = "\n".join(filter(None, [current_cells[column], value]))
        current_raw_rows.append(row.copy())
        current_end = physical_index
    flush()

    blocks: list[ExtractedBlock] = []
    for start, end, cells, raw_rows in logical_rows:
        fields = {keys[index]: value for index, value in enumerate(cells) if value}
        readable = "\n".join(
            f"{headers[index] or keys[index].replace('_', ' ').title()}: {value}"
            for index, value in enumerate(cells)
            if value
        )
        entity_name = cells[entity_column] if entity_column is not None else None
        metadata: dict[str, Any] = {
            "table_index": table_index,
            "section_path": section_path or section_title,
            "table_title": table_title,
            "headers": headers,
            "fields": fields,
            "table_row_from": start,
            "table_row_to": end,
            "zero_cross_row_overlap": True,
            "logical_row_group": f"{page_number or sheet_name or 'document'}:{table_index}:{start}-{end}",
        }
        if source_file:
            metadata["source_file"] = source_file
        if entity_name and entity_type:
            metadata.update({
                "entity_type": entity_type,
                "entity_name": entity_name,
                "entity_name_normalized": normalize_entity_name(entity_name),
            })
        raw = "\n".join(" | ".join(row) for row in raw_rows)
        blocks.append(ExtractedBlock(
            text=readable,
            raw=raw,
            content_type="table_row",
            extraction_method=extraction_method,
            page_from=page_number,
            page_to=page_number,
            sheet_name=sheet_name,
            row_from=start,
            row_to=end,
            section_title=section_title,
            metadata=metadata,
        ))
    return blocks


def text_outside_tables(page: Any, tables: list[Any]) -> str:
    if not tables:
        return page.extract_text(x_tolerance=2, y_tolerance=2) or ""

    def outside(obj: dict[str, Any]) -> bool:
        x0, x1 = float(obj.get("x0", 0)), float(obj.get("x1", 0))
        top, bottom = float(obj.get("top", 0)), float(obj.get("bottom", 0))
        center_x, center_y = (x0 + x1) / 2, (top + bottom) / 2
        return not any(
            left <= center_x <= right and table_top <= center_y <= table_bottom
            for left, table_top, right, table_bottom in (table.bbox for table in tables)
        )

    return page.filter(outside).extract_text(x_tolerance=2, y_tolerance=2) or ""


def ocr_pdf_page(path: Path, page_index: int) -> str:
    try:
        import pypdfium2 as pdfium
        import pytesseract

        pdf = pdfium.PdfDocument(str(path))
        try:
            image = pdf[page_index].render(scale=2.6).to_pil()
            return normalize_text(pytesseract.image_to_string(image, lang="eng+ara"))
        finally:
            pdf.close()
    except Exception as error:
        raise RuntimeError(
            f"Page {page_index + 1} needs OCR, but local OCR failed. "
            "Install Tesseract with English and Arabic language packs."
        ) from error


def pdf_page_is_visually_blank(path: Path, page_index: int) -> bool:
    """Distinguish a genuinely blank page from an image-only page needing OCR."""
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(str(path))
        try:
            image = pdf[page_index].render(scale=1).to_pil().convert("L")
            pixels = image.histogram()
            non_white = sum(pixels[:245])
            return non_white / max(1, image.width * image.height) < 0.0005
        finally:
            pdf.close()
    except Exception:
        return False


def extract_pdf(path: Path) -> list[ExtractedBlock]:
    blocks: list[ExtractedBlock] = []
    current_heading: str | None = None
    with pdfplumber.open(path) as document:
        for page_index, page in enumerate(document.pages):
            page_number = page_index + 1
            tables = page.find_tables()
            complete_raw_text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            raw_text = text_outside_tables(page, tables)
            method = "native"
            if (
                not tables
                and len(re.sub(r"\W", "", complete_raw_text, flags=re.UNICODE)) < 25
                and pdf_page_is_visually_blank(path, page_index)
            ):
                continue
            if not tables and len(re.sub(r"\W", "", complete_raw_text, flags=re.UNICODE)) < 25:
                raw_text = ocr_pdf_page(path, page_index)
                method = "ocr"
            text = normalize_text(raw_text)
            if not text and not tables:
                raise RuntimeError(f"PDF page {page_number} produced no usable text.")

            lines = [line.strip() for line in text.splitlines() if line.strip()]
            if lines and len(lines[0]) < 110 and (lines[0].endswith(":") or lines[0].istitle()):
                current_heading = lines[0].rstrip(":")
            if text:
                blocks.append(
                    ExtractedBlock(
                        text=text,
                        raw=raw_text,
                        content_type="section",
                        extraction_method=method,
                        page_from=page_number,
                        page_to=page_number,
                        section_title=current_heading,
                        metadata={
                            "source_page": page_number,
                            "source_file": path.name,
                            "section_path": current_heading,
                            "heading_level": 1,
                        },
                    )
                )
            for table_index, table in enumerate(tables):
                blocks.extend(table_row_blocks(
                    table.extract(),
                    extraction_method=method,
                    section_title=current_heading,
                    section_path=current_heading,
                    table_index=table_index,
                    page_number=page_number,
                    source_file=path.name,
                ))
    return blocks


def extract_docx(path: Path) -> list[ExtractedBlock]:
    document = Document(path)
    blocks: list[ExtractedBlock] = []
    heading: str | None = None
    heading_path: str | None = None
    heading_level = 1
    heading_stack: dict[int, str] = {}
    paragraph_number = 0
    for paragraph in document.paragraphs:
        paragraph_number += 1
        raw = paragraph.text
        text = normalize_text(raw)
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level_match = re.search(r"(\d+)", style)
            heading_level = int(level_match.group(1)) if level_match else 1
            heading_stack = {level: value for level, value in heading_stack.items() if level < heading_level}
            heading_stack[heading_level] = text
            heading = text
            heading_path = " > ".join(heading_stack[level] for level in sorted(heading_stack))
            content_type = "section"
        elif style.startswith("list"):
            content_type = "list"
        else:
            content_type = "paragraph"
        blocks.append(
            ExtractedBlock(
                text=text,
                raw=raw,
                content_type=content_type,
                extraction_method="docx",
                section_title=heading,
                metadata={
                    "paragraph": paragraph_number,
                    "style": paragraph.style.name,
                    "section_path": heading_path,
                    "heading_level": heading_level,
                },
            )
        )
    for table_index, table in enumerate(document.tables):
        blocks.extend(table_row_blocks(
            [[cell.text for cell in row.cells] for row in table.rows],
            extraction_method="docx",
            section_title=heading,
            section_path=heading_path,
            table_index=table_index,
            source_file=path.name,
        ))
    return blocks


def extract_xlsx(path: Path) -> list[ExtractedBlock]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[ExtractedBlock] = []
    try:
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            header_index = next((i for i, row in enumerate(rows) if any(value is not None for value in row)), None)
            if header_index is None:
                continue
            blocks.extend(table_row_blocks(
                [[str(value) if value is not None else "" for value in row] for row in rows[header_index:]],
                extraction_method="xlsx",
                section_title=sheet.title,
                section_path=sheet.title,
                table_index=0,
                sheet_name=sheet.title,
                source_file=path.name,
            ))
    finally:
        workbook.close()
    return blocks


def extract_xlsb(path: Path) -> list[ExtractedBlock]:
    """Extract binary Excel workbooks while preserving sheet and row sources."""
    try:
        from pyxlsb import open_workbook
    except ImportError as error:
        raise RuntimeError(
            "XLSB extraction requires pyxlsb. Install insurance_ingestion/requirements.txt."
        ) from error

    blocks: list[ExtractedBlock] = []
    with open_workbook(str(path)) as workbook:
        for sheet_name in workbook.sheets:
            with workbook.get_sheet(sheet_name) as sheet:
                rows = [
                    [normalize_text(str(cell.v)) if cell.v is not None else "" for cell in row]
                    for row in sheet.rows()
                ]
            header_index = next((index for index, row in enumerate(rows) if any(row)), None)
            if header_index is None:
                continue
            blocks.extend(table_row_blocks(
                rows[header_index:],
                extraction_method="xlsb",
                section_title=sheet_name,
                section_path=sheet_name,
                table_index=0,
                sheet_name=sheet_name,
                source_file=path.name,
            ))
    return blocks


def sliding_chunks(words: list[str], maximum: int = 340, overlap: int = 40) -> Iterable[list[str]]:
    start = 0
    while start < len(words):
        end = min(len(words), start + maximum)
        yield words[start:end]
        if end == len(words):
            break
        start = max(start + 1, end - overlap)


def chunk_blocks(blocks: list[ExtractedBlock]) -> list[ExtractedBlock]:
    chunks: list[ExtractedBlock] = []
    seen: set[str] = set()
    for block in blocks:
        # Structured records are semantic units. Never merge, overlap, or split them
        # across an adjacent row merely to approach a target token size.
        if block.content_type in {"table", "table_row", "list"} or token_estimate(block.text) <= 380:
            candidates = [block]
        else:
            words = block.text.split()
            candidates = [
                ExtractedBlock(
                    text=" ".join(part), raw=block.raw, content_type=block.content_type,
                    extraction_method=block.extraction_method, page_from=block.page_from,
                    page_to=block.page_to, sheet_name=block.sheet_name, row_from=block.row_from,
                    row_to=block.row_to, section_title=block.section_title,
                    metadata={**block.metadata, "split": True},
                )
                for part in sliding_chunks(words)
            ]
        for candidate in candidates:
            digest = content_hash(candidate.text)
            if digest not in seen:
                seen.add(digest)
                chunks.append(candidate)
    return chunks


class LocalEmbeddingProvider:
    dimension = 384
    index_model = "gte-small"

    def __init__(self, model_name: str | None = None) -> None:
        from sentence_transformers import SentenceTransformer

        # Query-time embeddings are generated by Supabase.ai gte-small. Stored
        # passages must use the same model family, pooling and normalization;
        # vectors from different models are not comparable.
        self.model_name = model_name or os.environ.get(
            "INSURANCE_EMBEDDING_MODEL",
            "thenlper/gte-small",
        )
        self.model = SentenceTransformer(self.model_name)

    def embed_passages(self, values: list[str]) -> list[list[float]]:
        vectors = self.model.encode(
            values,
            normalize_embeddings=True,
            show_progress_bar=len(values) > 32,
        )
        return [vector.tolist() for vector in vectors]


class InsuranceIngestionWorker:
    bucket = "insurance-documents"
    worker_version = "2026.08.13-hardening-v2"

    def __init__(self, client: Client, skip_embeddings: bool = False) -> None:
        self.client = client
        self.skip_embeddings = skip_embeddings
        self.embedding_provider = None if skip_embeddings else LocalEmbeddingProvider()

    def register(self, path: Path, title: str, version: str | None, category: str | None) -> str:
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED:
            raise ValueError(f"Unsupported file type: {suffix}. Expected PDF, DOCX, or XLSX.")
        payload = path.read_bytes()
        checksum = sha256_bytes(payload)
        existing = self.client.table("insurance_documents").select("id").eq("checksum", checksum).execute().data
        if existing:
            return str(existing[0]["id"])
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "-", path.name).strip("-")
        storage_path = f"{checksum[:12]}/{safe_name}"
        uploaders = (
            self.client.table("app_users")
            .select("user_id")
            .eq("role", "inventory")
            .eq("is_active", True)
            .limit(1)
            .execute()
            .data
        )
        if not uploaders:
            raise RuntimeError("No active inventory user is available for document audit ownership.")
        uploaded_by = uploaders[0]["user_id"]
        self.client.storage.from_(self.bucket).upload(
            storage_path,
            payload,
            {
                "content-type": MIME_TYPES.get(
                    suffix,
                    mimetypes.guess_type(path.name)[0] or "application/octet-stream",
                ),
                "upsert": "true",
            },
        )
        row = self.client.table("insurance_documents").insert(
            {
                "file_name": safe_name,
                "original_file_name": path.name,
                "storage_path": storage_path,
                "mime_type": MIME_TYPES[suffix],
                "file_extension": suffix.lstrip("."),
                "file_size": len(payload),
                "title": title,
                "uploaded_by": uploaded_by,
                "version": version,
                "document_category": category,
                "checksum": checksum,
                "processing_status": "queued",
                "embedding_model": self.embedding_provider.index_model
                if self.embedding_provider else None,
                "metadata": {
                    "ingested_from": "local_worker",
                    "embedding_model": self.embedding_provider.index_model
                    if self.embedding_provider else None,
                },
            }
        ).execute().data[0]
        self.client.table("insurance_ingestion_jobs").insert({"document_id": row["id"]}).execute()
        return str(row["id"])

    def process(self, document_id: str, local_path: Path | None = None) -> None:
        document = self.client.table("insurance_documents").select("*").eq("id", document_id).single().execute().data
        reprocessing_ready_document = document.get("processing_status") == "ready"
        if reprocessing_ready_document:
            self.client.table("insurance_documents").update(
                {"processing_error": None, "extraction_started_at": utc_now()}
            ).eq("id", document_id).execute()
        else:
            self._status(document_id, "extracting", extraction_started_at=utc_now())
        path = local_path
        temporary = False
        if path is None:
            import tempfile

            suffix = f".{document['file_extension']}"
            handle = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            handle.write(self.client.storage.from_(self.bucket).download(document["storage_path"]))
            handle.close()
            path = Path(handle.name)
            temporary = True
        try:
            extractors = {
                ".pdf": extract_pdf,
                ".docx": extract_docx,
                ".xlsx": extract_xlsx,
                ".xlsb": extract_xlsb,
            }
            blocks = extractors[path.suffix.lower()](path)
            if not reprocessing_ready_document:
                self._status(document_id, "chunking")
            chunks = chunk_blocks(blocks)
            if not chunks:
                raise RuntimeError("Extraction completed but produced no searchable chunks.")
            embeddings: list[list[float] | None]
            if self.embedding_provider:
                if not reprocessing_ready_document:
                    self._status(document_id, "embedding")
                embeddings = self.embedding_provider.embed_passages([chunk.text for chunk in chunks])
            else:
                embeddings = [None] * len(chunks)

            rows = []
            for index, (chunk, embedding) in enumerate(zip(chunks, embeddings, strict=True)):
                document_topic = normalize_entity_name(
                    str(document.get("document_category") or document.get("title") or "")
                )
                rows.append(
                    {
                        "document_id": document_id,
                        "chunk_index": index,
                        "page_from": chunk.page_from,
                        "page_to": chunk.page_to,
                        "sheet_name": chunk.sheet_name,
                        "row_from": chunk.row_from,
                        "row_to": chunk.row_to,
                        "section_title": chunk.section_title,
                        "content_text": chunk.text,
                        "raw_content": chunk.raw,
                        "content_type": chunk.content_type,
                        "extraction_method": chunk.extraction_method,
                        "token_count": token_estimate(chunk.text),
                        "content_hash": content_hash(chunk.text),
                        "embedding": embedding,
                        "embedding_model": self.embedding_provider.index_model
                        if self.embedding_provider else None,
                        "metadata": {
                            **chunk.metadata,
                            "document_topic": document.get("title"),
                            "topic_normalized": document_topic,
                            "document_category": document.get("document_category"),
                            "document_version": document.get("version"),
                            "parent_group": chunk.metadata.get("logical_row_group")
                            or f"{chunk.page_from or chunk.sheet_name or 'document'}:{chunk.section_title or 'root'}",
                            "embedding_model": self.embedding_provider.index_model
                            if self.embedding_provider else None,
                        },
                        "parent_group": chunk.metadata.get("logical_row_group")
                        or f"{chunk.page_from or chunk.sheet_name or 'document'}:{chunk.section_title or 'root'}",
                        "topic": document.get("title"),
                        "topic_normalized": document_topic,
                        "structured_fields": chunk.metadata.get("fields") or {},
                        "numeric_facts": [],
                    }
                )
            processing_run_id = str(uuid.uuid4())
            replacement = self.client.rpc(
                "replace_insurance_document_chunks_v2",
                {
                    "p_document_id": document_id,
                    "p_processing_run_id": processing_run_id,
                    "p_rows": rows,
                    "p_worker_version": self.worker_version,
                },
            ).execute().data
            if not replacement or int(replacement.get("chunk_count", 0)) != len(rows):
                raise RuntimeError("Atomic chunk replacement did not persist the complete generation.")
            aliases = extract_entity_aliases(
                chunks,
                document_id=document_id,
                document_title=str(document["title"]),
            )
            profile = document_profile(chunks, title=str(document["title"]))
            catalog = extract_document_catalog(
                chunks,
                document_id=document_id,
                document_title=str(document["title"]),
                aliases=aliases,
            )
            health = document_health_report(chunks, profile, catalog)
            if health["status"] != "verified":
                raise RuntimeError(
                    "Document health validation failed: "
                    + ", ".join(name for name, passed in health["checks"].items() if not passed)
                )
            self.client.table("insurance_document_entities").delete().eq(
                "document_id", document_id
            ).execute()
            if aliases:
                self.client.table("insurance_entity_aliases").upsert(
                    aliases,
                    on_conflict="entity_type,normalized_alias",
                ).execute()
                document_entities = [
                    {
                        "document_id": document_id,
                        "entity_type": alias["entity_type"],
                        "canonical_name": alias["canonical_name"],
                        "normalized_entity": alias["normalized_alias"],
                        "role": "class" if alias["entity_type"] == "therapy_class" else "primary",
                        "confidence": 1,
                        "metadata": {"source": alias["metadata"].get("discovery")},
                    }
                    for alias in aliases
                ]
                self.client.table("insurance_document_entities").upsert(
                    document_entities,
                    on_conflict="document_id,entity_type,normalized_entity,role",
                ).execute()
            # The separate catalog/profile tables retain generic document
            # structure for future policies without application-code changes.
            self.client.table("insurance_document_profiles").upsert(
                {
                    "document_id": document_id,
                    "profile": profile,
                    "status": "verified",
                    "extraction_version": self.worker_version,
                },
                on_conflict="document_id",
            ).execute()
            self.client.table("insurance_document_sections").delete().eq(
                "document_id", document_id
            ).execute()
            section_rows = [
                {
                    "document_id": document_id,
                    "section_path": str(section["path"]),
                    "title": str(section["title"]),
                    "normalized_title": normalize_entity_name(str(section["title"])),
                    "section_order": index,
                    "metadata": {
                        "source": "extracted_heading",
                        "heading_level": int(section.get("level") or 1),
                    },
                }
                for index, section in enumerate(profile["sections"])
            ]
            if section_rows:
                self.client.table("insurance_document_sections").upsert(
                    section_rows,
                    on_conflict="document_id,section_path",
                ).execute()
            self.client.table("insurance_entity_catalog").delete().eq(
                "document_id", document_id
            ).execute()
            if catalog:
                self.client.table("insurance_entity_catalog").upsert(
                    catalog,
                    on_conflict="document_id,entity_type,normalized_entity",
                ).execute()
            self.client.table("insurance_document_health_checks").insert(
                {
                    "document_id": document_id,
                    "run_id": processing_run_id,
                    "status": health["status"],
                    "report": health,
                    "worker_version": self.worker_version,
                }
            ).execute()
            final_status = "ready" if self.embedding_provider else "embedding"
            validation_status = "verified" if self.embedding_provider else "pending"
            self.client.table("insurance_documents").update(
                {
                    "processing_status": final_status,
                    "processing_error": None,
                    "extraction_completed_at": utc_now(),
                    "embedding_model": self.embedding_provider.index_model
                    if self.embedding_provider else None,
                    "search_validation_status": validation_status,
                    "search_validated_at": utc_now() if self.embedding_provider else None,
                    "metadata": {
                        **(document.get("metadata") or {}),
                        "chunk_count": len(rows),
                        "document_profile": profile,
                        "document_health": health,
                        "embedding_model": self.embedding_provider.index_model
                        if self.embedding_provider else None,
                    },
                }
            ).eq("id", document_id).execute()
            self.client.table("insurance_processing_runs").update(
                {
                    "status": "ready" if self.embedding_provider else "embedding",
                    "validation_status": validation_status,
                    "validation_report": {
                        "chunk_count": len(rows),
                        "entity_count": len(aliases),
                        "atomic_replacement": True,
                    },
                    "completed_at": utc_now() if self.embedding_provider else None,
                }
            ).eq("id", processing_run_id).execute()
            self.client.table("insurance_ingestion_jobs").update(
                {
                    "status": "completed" if self.embedding_provider else "running",
                    "updated_at": utc_now(),
                    "last_error": None,
                }
            ).eq("document_id", document_id).execute()
        except Exception as error:
            self.client.table("insurance_documents").update(
                {"processing_status": "failed", "processing_error": str(error)[:2000]}
            ).eq("id", document_id).execute()
            self.client.table("insurance_ingestion_jobs").update(
                {"status": "failed", "updated_at": utc_now(), "last_error": str(error)[:2000]}
            ).eq("document_id", document_id).execute()
            raise
        finally:
            if temporary and path:
                path.unlink(missing_ok=True)

    def _status(self, document_id: str, status: str, **extra: Any) -> None:
        self.client.table("insurance_documents").update(
            {"processing_status": status, "processing_error": None, **extra}
        ).eq("id", document_id).execute()


def build_client() -> Client:
    from supabase import create_client

    url = os.environ.get("SUPABASE_URL", "").strip()
    key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "").strip()
    if not url or not key:
        raise RuntimeError("Set SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY for the worker process only.")
    return create_client(url, key)


def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest private insurance guideline documents into Supabase.")
    parser.add_argument("path", type=Path)
    parser.add_argument("--title")
    parser.add_argument("--version")
    parser.add_argument("--category")
    parser.add_argument("--skip-embeddings", action="store_true")
    args = parser.parse_args()
    path = args.path.expanduser().resolve()
    if not path.is_file():
        raise FileNotFoundError(path)
    worker = InsuranceIngestionWorker(build_client(), skip_embeddings=args.skip_embeddings)
    document_id = worker.register(path, args.title or path.stem, args.version, args.category)
    worker.process(document_id, path)
    print(f"Ready: {document_id} ({path.name}) on {socket.gethostname()}")


if __name__ == "__main__":
    main()
