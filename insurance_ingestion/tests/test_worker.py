from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen.canvas import Canvas

from insurance_ingestion.worker import (
    ExtractedBlock,
    chunk_blocks,
    document_health_report,
    document_profile,
    extract_document_catalog,
    extract_docx,
    extract_entity_aliases,
    extract_pdf,
    extract_xlsx,
    normalize_text,
    table_row_blocks,
)


class InsuranceIngestionTests(unittest.TestCase):
    def test_normalization_preserves_medical_values(self) -> None:
        value = normalize_text("Ubrogepant   200 mg\n\n\nAge ≥ 18 years")
        self.assertEqual(value, "Ubrogepant 200 mg\n\nAge ≥ 18 years")

    def test_chunking_is_deduplicated_and_bounded(self) -> None:
        long_text = " ".join(f"clinical{i}" for i in range(900))
        chunks = chunk_blocks([
            ExtractedBlock(text=long_text, raw=long_text),
            ExtractedBlock(text=long_text, raw=long_text),
        ])
        self.assertGreater(len(chunks), 1)
        self.assertEqual(len({chunk.text for chunk in chunks}), len(chunks))
        self.assertTrue(all(len(chunk.text.split()) <= 340 for chunk in chunks))

    def test_pdf_extraction_preserves_page_number(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "guideline.pdf"
            canvas = Canvas(str(path))
            canvas.drawString(72, 760, "CGRP Coverage Guideline")
            canvas.drawString(72, 735, "Ubrogepant maximum dose is 200 mg in 24 hours.")
            canvas.save()
            blocks = extract_pdf(path)
            self.assertEqual(blocks[0].page_from, 1)
            self.assertIn("200 mg", blocks[0].text)

    def test_pdf_extraction_skips_a_truly_blank_page_without_ocr(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "guideline-with-blank-page.pdf"
            canvas = Canvas(str(path))
            canvas.drawString(72, 760, "PPI coverage requires a documented diagnosis.")
            canvas.showPage()
            canvas.showPage()
            canvas.save()
            blocks = extract_pdf(path)
            self.assertEqual([block.page_from for block in blocks], [1])

    def test_docx_extraction_preserves_heading_and_table(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "guideline.docx"
            document = Document()
            document.add_heading("Coverage Criteria", level=1)
            document.add_paragraph("Prior therapy is required.")
            document.add_heading("Initial approval", level=2)
            document.add_paragraph("A current report is required.")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Medication"
            table.cell(0, 1).text = "Dose"
            table.cell(1, 0).text = "Erenumab"
            table.cell(1, 1).text = "70 mg monthly"
            document.save(path)
            blocks = extract_docx(path)
            self.assertTrue(any(block.section_title == "Coverage Criteria" for block in blocks))
            self.assertTrue(any(block.content_type == "table_row" and "Erenumab" in block.text for block in blocks))
            nested = next(block for block in blocks if block.text == "A current report is required.")
            self.assertEqual(nested.metadata["section_path"], "Coverage Criteria > Initial approval")
            self.assertEqual(nested.metadata["heading_level"], 2)

    def test_xlsx_extraction_preserves_sheet_and_row(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "guideline.xlsx"
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Migraine"
            sheet.append(["Medication", "Dose"])
            sheet.append(["Rimegepant", "75 mg every other day"])
            workbook.save(path)
            blocks = extract_xlsx(path)
            self.assertEqual(blocks[0].sheet_name, "Migraine")
            self.assertEqual(blocks[0].row_from, 2)
            self.assertIn("Rimegepant", blocks[0].text)
            self.assertEqual(blocks[0].metadata["entity_name"], "Rimegepant")
            self.assertEqual(blocks[0].metadata["fields"]["dose"], "75 mg every other day")

    def test_medication_table_rows_never_cross_entity_boundaries(self) -> None:
        table = [
            ["Medication", "Indications", "Recommended Dose", "Route"],
            ["Ubrogepant", "Acute treatment of migraine", "50 mg or 100 mg orally; max 200 mg in 24 hours", "Oral"],
            ["", "Not used as preventive", "May repeat after 2 hours", ""],
            ["Atogepant", "Preventive treatment of episodic migraine", "10 mg, 30 mg, or 60 mg once daily", "Oral"],
            ["Rimegepant", "Acute and preventive treatment", "75 mg every other day", "Oral"],
        ]
        chunks = chunk_blocks(table_row_blocks(
            table,
            extraction_method="native",
            section_title="CGRP inhibitors",
            table_index=0,
            page_number=4,
            source_file="cgrp.pdf",
        ))

        self.assertEqual([chunk.metadata["entity_name"] for chunk in chunks], [
            "Ubrogepant", "Atogepant", "Rimegepant",
        ])
        ubrogepant, atogepant, rimegepant = chunks
        self.assertIn("max 200 mg in 24 hours", ubrogepant.text)
        self.assertIn("May repeat after 2 hours", ubrogepant.text)
        self.assertNotIn("Atogepant", ubrogepant.text)
        self.assertNotIn("10 mg, 30 mg, or 60 mg", ubrogepant.text)
        self.assertNotIn("Ubrogepant", atogepant.text)
        self.assertNotIn("Rimegepant", atogepant.text)
        self.assertNotIn("Atogepant", rimegepant.text)
        self.assertTrue(all(chunk.metadata["zero_cross_row_overlap"] for chunk in chunks))

    def test_entity_registry_includes_structured_and_prose_medications(self) -> None:
        chunks = [
            ExtractedBlock(
                text="Medication: Rimegepant\nDose: 75 mg",
                raw="Rimegepant | 75 mg",
                metadata={"entity_type": "medication", "entity_name": "Rimegepant"},
            ),
            ExtractedBlock(
                text="OZEMPIC 0.25 MG and MOUNJARO 2.5 MG are initial non-therapeutic doses. TRULICITY 0.75 MG needs a report.",
                raw="",
            ),
        ]
        aliases = extract_entity_aliases(
            chunks,
            document_id="glp-document",
            document_title="GLP-1 Receptor Agonists",
        )
        by_name = {
            item["canonical_name"]: item
            for item in aliases
            if item["entity_type"] == "medication"
        }
        therapy_classes = {
            item["canonical_name"]
            for item in aliases
            if item["entity_type"] == "therapy_class"
        }

        self.assertEqual(set(by_name), {"Rimegepant", "Ozempic", "Mounjaro", "Trulicity"})
        self.assertEqual(therapy_classes, {"Glp-1"})
        self.assertEqual(by_name["Mounjaro"]["metadata"]["strength"], "2.5 mg")
        self.assertEqual(by_name["Mounjaro"]["metadata"]["document_id"], "glp-document")

    def test_entity_registry_rejects_common_words_before_numeric_values(self) -> None:
        chunks = [ExtractedBlock(
            text="The 0.25 mg dose is used, then 2.4 mg may be considered.",
            raw="",
        )]
        aliases = extract_entity_aliases(
            chunks,
            document_id="future-document",
            document_title="Future policy",
        )
        self.assertEqual(aliases, [])

    def test_future_medication_is_cataloged_from_a_trusted_table_header(self) -> None:
        chunks = table_row_blocks(
            [["Medication", "Minimum age", "Required report"], ["Drug XYZ", "15 years", "Specialist report"]],
            extraction_method="native",
            section_title="Coverage criteria",
            table_index=0,
            page_number=1,
        )
        aliases = extract_entity_aliases(chunks, document_id="future", document_title="Future policy")
        catalog = extract_document_catalog(
            chunks, document_id="future", document_title="Future policy", aliases=aliases,
        )
        medicines = {item["canonical_name"] for item in catalog if item["entity_type"] == "medication"}
        self.assertIn("Drug XYZ", medicines)
        report = document_health_report(chunks, document_profile(chunks, title="Future policy"), catalog)
        self.assertEqual(report["status"], "verified")

    def test_single_column_wrapped_table_keeps_one_semantic_fact(self) -> None:
        blocks = table_row_blocks(
            [
                ["Icosapent Ethyl: 4 grams/day as two 1-gram"],
                ["capsules twice daily with food."],
            ],
            extraction_method="native",
            section_title="Dosing",
            table_index=0,
            page_number=2,
        )
        self.assertEqual(len(blocks), 1)
        self.assertIn("two 1-gram capsules twice daily with food", blocks[0].text)
        self.assertTrue(blocks[0].metadata["merged_wrapped_rows"])


if __name__ == "__main__":
    unittest.main()
